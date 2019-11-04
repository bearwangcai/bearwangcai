import os  
import docx
from docx import Document
from read_docx import read_file
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

global professions
global province_names #省名
global content_names #共性意见 个性意见
global all

def write_content(document, file, profession, province_name, line_spacing, space_after):
    '''
    写入具体评审意见，包含共性问题，个性问题、个性问题等内容
    '''
    for content_name in content_names: #共性问题，个性问题
        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = Pt(line_spacing) #行间距
        p.paragraph_format.space_after = Pt(space_after)
        if content_name == '（一）共性意见':
            p.add_run('（一）你省共性要求').bold = True
        else:
            p.add_run('（二）你省其他审查意见').bold = True
        for content in file[profession][province_name][content_name]:
            if not (content == ''):
                p = document.add_paragraph(content)
                p.paragraph_format.line_spacing = Pt(line_spacing) #行间距
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.first_line_indent = Pt(32) #首行缩进

def set_style(run, font_name = u'黑体', font_size = 18, bold = True):
    '''
    run
    font_name = u'黑体'
    font_size = 18  
    bold = True
    '''
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) #更改字体必须.font.name和._element.rPr.rFonts.set两个都要改
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x0a,0x0a,0x0a)
    run.bold = bold

def write_file(file):

    for province_name in province_names:
        document = Document()
        document.styles['Normal'].font.name = u'仿宋'
        document.styles['Normal'].font.size = Pt(16)
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        line_spacing = 32 #行间距
        space_after =  0 #段后间距
        
        #标题行
        p = document.add_paragraph('')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #只有add_paragraph方法有这个属性
        run = p.add_run(province_name + '分公司2020-2022年网络发展滚动规划') #只有add_paragraph.add_run这个方法有下面的属性，而且想要更改个别属性，就需要用.add_run方法！
        set_style(run, u'黑体', 18, True)

        p = document.add_paragraph('')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('初审评审意见')
        set_style(run, u'黑体', 18, True)
        
        
        p = document.add_paragraph('')
        p = document.add_paragraph('    2019年10月28日-10月30日集团公司网络发展部组织集团规划项目组对你省公司2020-2022年网络发展滚动规划总体思路进行了集中评审，形成审查意见如下。')
        p.paragraph_format.line_spacing = Pt(line_spacing)
        p.paragraph_format.space_after = Pt(space_after)
        p = document.add_paragraph('')


        #网络规划
        p = document.add_paragraph('')
        run = p.add_run('网络规划') 
        set_style(run, u'黑体', 18, False)

        # 一 总体要求
        run = document.add_heading('',level=1).add_run('一 总体要求')
        set_style(run, u'黑体', 16, True) 

        p = document.add_paragraph('    （一）细化梳理网络现状，充分利旧现网资源，提质增效。')
        p.paragraph_format.line_spacing = Pt(line_spacing)
        p.paragraph_format.space_after = Pt(space_after)
        p = document.add_paragraph('    （二）规划方案原则上必须有国家要求（当地政府要求）、集团战略、市场需求或技术发展趋势等作为规划输入。')
        p.paragraph_format.line_spacing = Pt(line_spacing)
        p.paragraph_format.space_after = Pt(space_after)
        p = document.add_paragraph('    （三）为保证2020年的5G建设投资，除5G以外的各专业投资预计将大幅压缩，')
        p.paragraph_format.line_spacing = Pt(line_spacing)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run('具体见各专业压缩比例。')
        #run.font.color.rgb = RGBColor(0xff,0x00,0x00) #想要个别更改属性，就需要用.add_run方法
        run1 = p.add_run('可建可不建的项目原则上不予建设。')


        # 二 各专业具体要求  
        p = document.add_paragraph('')
        run = document.add_heading('',level=1).add_run('二 各专业具体要求')
        set_style(run, u'黑体', 16, True)     

        for index_1, profession in enumerate(professions):
            
            #document.add_heading(profession, level=1)

            try:
                if (len(profession) == 1):
                    profession = profession[0]

                    #添加大专业标题
                    run = document.add_heading('',level=2).add_run('2.' + str(index_1+1) + ' ' + profession)
                    set_style(run, u'宋体', 16, True)
                    run.font.color.rgb = RGBColor(0x1a,0x1a,0x1a)
                    write_content(document, file, profession, province_name, line_spacing, space_after)
                else:
                    for index_2, sub_profession in enumerate(profession):

                        if index_2 == 0:
                            #添加大专业标题
                            run = document.add_heading('',level=2).add_run('2.' + str(index_1+1) + ' ' + sub_profession)
                            set_style(run, u'宋体', 16, True)
                            run.font.color.rgb = RGBColor(0x1a,0x1a,0x1a)

                        else:
                            #添加小专业标题
                            run = document.add_heading('',level=3).add_run('2.' + str(index_1+1) + '.' + str(index_2) + ' ' + sub_profession)
                            set_style(run, u'宋体 (中文标题)', 16, True)
                            run.font.color.rgb = RGBColor(0x1a,0x1a,0x1a)
                            write_content(document, file, sub_profession, province_name, line_spacing, space_after)
            except: 
                print(profession)
        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = Pt(line_spacing) #行间距
        p.paragraph_format.space_after = Pt(space_after)
        '''
        #节能减排
        p = document.add_paragraph('')
        run = p.add_run('节能减排')
        set_style(run, u'黑体', 18, False)
        write_content(document, file, '节能减排', province_name, line_spacing, space_after)
        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = Pt(line_spacing) #行间距
        p.paragraph_format.space_after = Pt(space_after)
        '''
        #应急通信
        p = document.add_paragraph('')
        run = p.add_run('应急通信')
        set_style(run, u'黑体', 18, False)
        write_content(document, file, '应急通信', province_name, line_spacing, space_after)


        document.save(r'C:\Users\x270\Desktop\test' + '\\'+ province_name + '分公司2020-2022年网络发展滚动规划初审意见.docx')



def main():
    global professions
    #professions = ['IDC',['无线网','移动网','核心网']]
    
    #professions = [['数据网'],['STN(IPRAN)'],['传输网'],['接入网'],['IDC'],['DC'],['云'],['业务平台'],['CDN'],['移动核心网'],['固网核心网'],['应急通信'],['节能减排']]

    professions = [['移动网','核心网'],['基础网','接入网和综合业务接入区','IP网','STN(IPRAN)','传输网'],['IDC及基础设施建设','IDC及DC基础设施建设','管道及其他基础设施'],['云计算','云','CDN','业务平台']]

    global province_names #省名
    province_names = ['北京','天津','河北','山西','内蒙古','辽宁','吉林','黑龙江','上海','江苏','浙江','安徽','福建','江西','山东','河南','湖北','湖南','广东','广西','海南','重庆','四川','贵州','云南','西藏','陕西','甘肃','青海','宁夏','新疆']
    global content_names #共性意见 个性意见
    content_names = ['（一）共性意见','（二）你省审查意见']
    global all
    all = read_file()
    write_file(all)

if __name__ == "__main__":
    main()




        




        
